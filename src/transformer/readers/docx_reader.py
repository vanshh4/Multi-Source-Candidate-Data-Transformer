"""
DOCX resume reader for Phase 2.

Extracts paragraph and table text from .docx files using python-docx, while
returning controlled statuses for missing, empty, corrupted, or unreadable files.
"""

from __future__ import annotations

from typing import Dict, List

from docx import Document

from transformer.readers.base import (
    STATUS_EMPTY,
    STATUS_MISSING,
    STATUS_OK,
    STATUS_UNREADABLE,
    build_reader_result,
)
from transformer.utils.errors import (
    DOCX_NO_TEXT,
    SOURCE_UNREADABLE,
    empty_source_warning,
    format_exception_message,
    missing_source_warning,
)
from transformer.utils.file_utils import file_exists, get_file_name, is_empty_file


def _extract_table_text(document: Document) -> List[str]:
    """Extract cell text from all tables in a DOCX document."""
    table_lines: List[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))
    return table_lines


def read_docx_source(
    path: str,
    *,
    source_id: str = "resume_docx",
    source_type: str = "resume_docx",
    required: bool = False,
) -> Dict:
    """Safely extract text from a DOCX resume."""
    source_name = get_file_name(path)

    if not file_exists(path):
        return build_reader_result(
            source_id=source_id,
            source_type=source_type,
            source_name=source_name,
            path=path,
            required=required,
            status=STATUS_MISSING,
            content_type="text",
            content="",
            metadata={"paragraph_count": 0, "table_count": 0, "character_count": 0, "has_extractable_text": False},
            warnings=[missing_source_warning(path)],
            errors=[],
        )

    if is_empty_file(path):
        return build_reader_result(
            source_id=source_id,
            source_type=source_type,
            source_name=source_name,
            path=path,
            required=required,
            status=STATUS_EMPTY,
            content_type="text",
            content="",
            metadata={"paragraph_count": 0, "table_count": 0, "character_count": 0, "has_extractable_text": False},
            warnings=[empty_source_warning(path)],
            errors=[],
        )

    try:
        document = Document(path)
        paragraph_lines = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
        table_lines = _extract_table_text(document)
        all_lines = paragraph_lines + table_lines
        text = "\n".join(all_lines).strip()
    except Exception as exc:
        return build_reader_result(
            source_id=source_id,
            source_type=source_type,
            source_name=source_name,
            path=path,
            required=required,
            status=STATUS_UNREADABLE,
            content_type="text",
            content="",
            metadata={"paragraph_count": 0, "table_count": 0, "character_count": 0, "has_extractable_text": False},
            warnings=[],
            errors=[format_exception_message(SOURCE_UNREADABLE, exc)],
        )

    if not text:
        return build_reader_result(
            source_id=source_id,
            source_type=source_type,
            source_name=source_name,
            path=path,
            required=required,
            status=STATUS_EMPTY,
            content_type="text",
            content="",
            metadata={
                "paragraph_count": len(paragraph_lines),
                "table_count": len(document.tables),
                "character_count": 0,
                "has_extractable_text": False,
            },
            warnings=[DOCX_NO_TEXT],
            errors=[],
        )

    return build_reader_result(
        source_id=source_id,
        source_type=source_type,
        source_name=source_name,
        path=path,
        required=required,
        status=STATUS_OK,
        content_type="text",
        content=text,
        metadata={
            "paragraph_count": len(paragraph_lines),
            "table_count": len(document.tables),
            "character_count": len(text),
            "has_extractable_text": True,
        },
        warnings=[],
        errors=[],
    )


def read(path: str, **kwargs) -> Dict:
    """Friendly alias used by router/CLI."""
    return read_docx_source(path, **kwargs)
